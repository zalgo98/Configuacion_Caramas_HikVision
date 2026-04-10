"""
fill_document.py
================
Rellena la plantilla Word "Asignacion de IPs.docx" sustituyendo etiquetas
y guarda el resultado como "Asignacion de IPs - Torre <poste>.docx" en la
misma carpeta donde está la plantilla Word utilizada.

Etiquetas reconocidas
---------------------
{{ opera }}         → número de operación
{{ torre }}         → número de poste/torre
{{ ip_router }}     → IP del router (gateway)
{{ den_cam_X }}     → nombre de la cámara X  (X = 1…6, orden por posición)
{{ sn_cam_X }}      → número de serie de la cámara X (últimos 9 dígitos)
{{ ip_cam_X }}      → IP de la cámara X
{{ ubicación X }}   → nombre de la cámara asignada a esa posición del diagrama
                      (1=left_top, 2=left_mid, 3=left_bottom,
                       4=right_top, 5=right_mid, 6=right_bottom)
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# ── Orden canónico de posiciones → número de cámara ──────────────────────────
POSITION_ORDER = [
    "left_top",
    "left_mid",
    "left_bottom",
    "right_top",
    "right_mid",
    "right_bottom",
]


# ── Helpers ───────────────────────────────────────────────────────────────────

def _last9(serial: str) -> str:
    """Devuelve los últimos 9 caracteres del número de serie."""
    s = (serial or "").strip()
    return s[-9:] if len(s) >= 9 else s


def _ip_sort_key(ip: str):
    """Devuelve una clave de ordenación natural para IPv4."""
    try:
        return tuple(int(part) for part in (ip or "").strip().split("."))
    except Exception:
        return (999, 999, 999, 999)


def _build_replacements(
    operacion: str,
    poste: str,
    ip_router: str,
    assignments_info: dict,
) -> dict:
    """
    Construye el diccionario {etiqueta: valor} para todos los campos.

    assignments_info: {position_key: {"ip": ..., "serial": ..., "name": ...}}

    Reglas:
    - den_cam_X / dem_cam_X / sn_cam_X / ip_cam_X se rellenan por ORDEN DE IP.
    - ubicación X se rellena por la posición del diagrama.
    """
    replacements = {
        # Variantes con y sin espacios (el Word fragmenta los runs de formas distintas)
        "{{opera}}":       str(operacion),
        "{{ opera }}":     str(operacion),
        "{{ opera}}":      str(operacion),
        "{{opera }}":      str(operacion),
        "{{torre}}":       str(poste),
        "{{ torre }}":     str(poste),
        "{{ torre}}":      str(poste),
        "{{torre }}":      str(poste),
        "{{ip_router}}":   str(ip_router),
        "{{ ip_router }}": str(ip_router),
    }

    # 1) Tabla del documento: numeración independiente del diagrama, ordenada por IP
    sorted_cameras = sorted(
        (info for info in assignments_info.values() if info.get("ip")),
        key=lambda info: _ip_sort_key(info.get("ip", "")),
    )

    for idx in range(1, 7):
        info = sorted_cameras[idx - 1] if idx <= len(sorted_cameras) else {}
        name = info.get("name", "")
        serial = info.get("serial", "")
        ip = info.get("ip", "")

        for tag, value in [
            (f"{{{{den_cam_{idx}}}}}", name),
            (f"{{{{ den_cam_{idx} }}}}", name),
            (f"{{{{dem_cam_{idx}}}}}", name),
            (f"{{{{ dem_cam_{idx} }}}}", name),
            (f"{{{{sn_cam_{idx}}}}}", _last9(serial)),
            (f"{{{{ sn_cam_{idx} }}}}", _last9(serial)),
            (f"{{{{ip_cam_{idx}}}}}", ip),
            (f"{{{{ ip_cam_{idx} }}}}", ip),
        ]:
            replacements[tag] = value

    # 2) Diagrama: sigue la posición visual original
    for idx, pos_key in enumerate(POSITION_ORDER, start=1):
        info = assignments_info.get(pos_key, {})
        name = info.get("name", "")
        for tag, value in [
            (f"{{{{ ubicación {idx} }}}}", name),
            (f"{{{{ubicación {idx}}}}}", name),
        ]:
            replacements[tag] = value

    return replacements


# ── Sustitución en el XML de cada párrafo ────────────────────────────────────

def _replace_in_paragraph(paragraph, replacements: dict) -> None:
    """
    Sustituye todas las etiquetas en un párrafo.

    Las etiquetas pueden estar fragmentadas en varios <w:r> (runs).
    La estrategia:
      1. Reconstruir el texto completo del párrafo.
      2. Aplicar sustituciones sobre ese texto.
      3. Volcar el texto resultante en el primer run y vaciar el resto.
    """
    # Texto completo del párrafo
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    # ¿Contiene alguna etiqueta?
    modified = full_text
    for tag, value in replacements.items():
        modified = modified.replace(tag, value)

    if modified == full_text:
        return  # nada que cambiar

    # Volcar en los runs: primer run recibe todo, el resto se vacía
    if paragraph.runs:
        paragraph.runs[0].text = modified
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_table(table, replacements: dict) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            # Tablas anidadas
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)


# ── Sustitución en cuadros de texto (anchors/drawing shapes) ────────────────

def _replace_in_textboxes(doc, replacements: dict) -> None:
    """
    Recorre todos los cuadros de texto flotantes (wp:anchor > wps:txbx)
    y aplica las sustituciones.

    Las etiquetas en los cuadros de texto suelen estar fragmentadas en
    múltiples <w:r>, por lo que se reconstruye el texto completo del
    párrafo antes de sustituir, igual que en _replace_in_paragraph.
    """
    import re
    from lxml import etree

    # Namespaces relevantes
    WPC = "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
    WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    body = doc.element.body

    # Buscar todos los elementos wps:txbx en el documento
    txbx_elements = body.findall(f".//{{{WPS}}}txbx")

    for txbx in txbx_elements:
        # Cada txbx contiene un w:txbxContent con párrafos normales
        for p_elem in txbx.findall(f".//{{{W}}}p"):
            # Reconstruir texto de todos los runs
            runs = p_elem.findall(f".//{{{W}}}r")
            if not runs:
                continue

            full_text = "".join(
                (r.findtext(f"{{{W}}}t") or "") for r in runs
            )
            if not full_text.strip():
                continue

            modified = full_text
            for tag, value in replacements.items():
                modified = modified.replace(tag, value)

            if modified == full_text:
                continue

            # Volcar en el primer run, limpiar todos los demás
            first_t = runs[0].find(f"{{{W}}}t")
            if first_t is None:
                first_t = etree.SubElement(runs[0], f"{{{W}}}t")
            first_t.text = modified
            # Preservar espacios si el valor tiene espacios al inicio/fin
            if modified != modified.strip():
                first_t.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
            # Vaciar TODOS los runs siguientes (pueden tener o no <w:t>)
            for run in runs[1:]:
                t_elem = run.find(f"{{{W}}}t")
                if t_elem is not None:
                    t_elem.text = ""
                else:
                    # Añadir <w:t> vacío para evitar que queden fragmentos
                    etree.SubElement(run, f"{{{W}}}t").text = ""


# ── Punto de entrada público ──────────────────────────────────────────────────

def fill_ip_document(
    operacion: str,
    poste: str,
    assignments_info: dict,
    ip_router: str = "",
    template_path: str | None = None,
) -> str:
    """
    Rellena la plantilla y guarda el documento resultante en la misma carpeta
    donde se encuentra la plantilla Word utilizada.

    Parámetros
    ----------
    operacion       : número/código de operación
    poste           : número de torre/poste
    assignments_info: {position_key: {"ip": str, "serial": str, "name": str}}
    ip_router       : IP del router/gateway (opcional, cadena vacía si no se dispone)
    template_path   : ruta explícita a la plantilla; si es None se busca
                      automáticamente junto al ejecutable o en el directorio actual

    Devuelve
    --------
    Ruta absoluta del documento generado.
    """
    # ── Localizar plantilla ──────────────────────────────────────────────────
    if template_path is None:
        # Buscar en el directorio del script y en el directorio de trabajo
        candidates = [
            Path(__file__).parent / "Asignacion_IPs_template_publica.docx",
            Path(__file__).parent / "Asignacion de IPs.docx",
            Path(__file__).parent / "Asignación de IPs.docx",
            Path.cwd() / "Asignacion_IPs_template_publica.docx",
            Path.cwd() / "Asignacion de IPs.docx",
            Path.cwd() / "Asignación de IPs.docx",
        ]
        template_path = next((str(p) for p in candidates if p.exists()), None)
        if template_path is None:
            raise FileNotFoundError(
                "No se encontró la plantilla 'Asignacion de IPs.docx'. "
                "Colócala en el mismo directorio que la aplicación."
            )

    # ── Cargar plantilla ─────────────────────────────────────────────────────
    doc = Document(template_path)

    # ── Construir mapa de sustituciones ─────────────────────────────────────
    replacements = _build_replacements(operacion, poste, ip_router, assignments_info)

    # ── Aplicar sustituciones ────────────────────────────────────────────────
    # Párrafos del cuerpo principal
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # Tablas del cuerpo principal
    for table in doc.tables:
        _replace_in_table(table, replacements)

    # Encabezados y pies de página
    for section in doc.sections:
        for header_footer in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
                for table in header_footer.tables:
                    _replace_in_table(table, replacements)

    # Cuadros de texto flotantes (diagrama de ubicación)
    _replace_in_textboxes(doc, replacements)

    # ── Determinar destino: misma carpeta que la plantilla ───────────────────
    template_dir = Path(template_path).resolve().parent
    output_filename = f"Asignacion de IPs - Torre {poste}.docx"
    output_path = template_dir / output_filename

    doc.save(str(output_path))
    return str(output_path)
